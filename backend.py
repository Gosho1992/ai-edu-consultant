# backend.py

from fastapi.security import APIKeyHeader
api_key_header = APIKeyHeader(name="X-API-KEY", auto_error=False)

# External helpers
from parse_uploaded_file import parse_uploaded_file  # ← ensure file is named parse_uploaded_file.py
from content_analyzer import analyze_url_content

import os
import json
import hashlib
import logging
from datetime import datetime
from typing import Dict, List, Optional
from fastapi import HTTPException, Security
from dotenv import load_dotenv
from openai import OpenAI
from sentence_transformers import SentenceTransformer
from diskcache import Cache
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
import pytesseract
from PIL import Image
import pdfplumber
import requests
from bs4 import BeautifulSoup
import feedparser
import io
import magic
from pathlib import Path
from docx import Document

# ---------- Tesseract path (so OCR works on both Windows/Linux) ----------
if os.name == "nt":  # Windows
    pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
else:  # Linux/Mac (Streamlit Cloud uses Linux container)
    pytesseract.pytesseract.tesseract_cmd = "/usr/bin/tesseract"

# ---------- Configuration ----------
load_dotenv()
API_KEY_NAME = "X-API-KEY"
api_key_header = APIKeyHeader(name=API_KEY_NAME, auto_error=False)

class SecurityManager:
    @staticmethod
    def validate_api_key(api_key: str = Security(api_key_header)):
        if api_key != os.getenv("API_KEY"):
            raise HTTPException(403, "Invalid API key")
        return True


class EducationAgent:
    def __init__(self):
        self.client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        self.embedding_model = SentenceTransformer("all-MiniLM-L6-v2")
        self.cache = Cache("edubot_cache")
        self.security = SecurityManager()
        self._init_user_profile()
        self._init_services()
        self._setup_monitoring()

    # ---------- Bootstrapping ----------
    def _init_user_profile(self):
        self.user = {
            "academic": {"degree": None, "field": None, "gpa": None, "target_countries": []},
            "financial": {"budget": None, "scholarship_needs": True},
            "documents": {"cv": None, "sop": None, "transcripts": None},
            "engagement": {"last_active": datetime.now(), "query_count": 0},
        }

    def _init_services(self):
        self.services = {
            "scholarship_feeds": [
                "https://scholarshipscorner.website/feed/",
                "https://scholarshipunion.com/feed/",
            ],
            "apis": {
                "universities": "http://universities.hipolabs.com/search",
                "ranking": "https://edurank.org/api/unis.json",
            },
            "university_db": [],
        }

    def _setup_monitoring(self):
        self.metrics = {"gpt_calls": 0, "cache_hits": 0, "response_times": []}
        logging.basicConfig(
            filename="edubot.log",
            level=logging.INFO,
            format="%(asctime)s - %(levelname)s - %(message)s",
        )

    # ---------- Utilities ----------
    def _get_mime_type(self, file_bytes: bytes) -> str:
        return magic.from_buffer(file_bytes, mime=True)

    def _validate_file(self, file_bytes: bytes) -> bool:
        if not file_bytes:
            return False
        if len(file_bytes) > 200 * 1024 * 1024:  # 200MB app limit
            return False
        return True

    def _detect_file_type(self, file_bytes: bytes, filename: str) -> str:
        """
        Returns normalized 'pdf' | 'docx' | 'jpg' | 'png' | 'txt' | fallback to extension.
        """
        try:
            mime = magic.from_buffer(file_bytes, mime=True)
            ext = Path(filename).suffix[1:].lower()
            type_map = {
                "application/pdf": "pdf",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
                "application/msword": "doc",  # will treat as unsupported
                "image/jpeg": "jpg",
                "image/jpg": "jpg",
                "image/png": "png",
                "text/plain": "txt",
            }
            return type_map.get(mime, ext if ext else "unknown")
        except Exception as e:
            logging.error(f"File type detection failed: {str(e)}")
            return Path(filename).suffix[1:].lower() or "unknown"

    def _is_supported(self, file_type: str, doc_type: str) -> bool:
        supported_types = {
            "sop": {"pdf", "docx", "jpg", "png"},
            "cv": {"pdf", "docx"},
            "transcript": {"pdf", "jpg", "png"},
        }
        return file_type in supported_types.get(doc_type.lower(), set())

    # ---------- URL analyzer (kept simple) ----------
    def analyze_url(self, url: str) -> Dict:
        try:
            text = analyze_url_content(url)
            if text.startswith("⚠️ Error"):
                return {"text": "", "source": url, "error": text}
            return {"text": text, "source": url, "error": ""}
        except Exception as e:
            return {"text": "", "source": url, "error": f"❌ URL analysis failed: {str(e)}"}

    # ---------- Document analyzer (single, authoritative) ----------
    def analyze_document(self, file_bytes: bytes, filename: str, doc_type: str, purpose: Optional[str] = None, extra_context: Optional[str] = None) -> Dict:
    try:
        logging.info(f"Analyzing document: {filename}, type: {doc_type}, purpose: {purpose}")

        # Validate file
        if not self._validate_file(file_bytes):
            return {
                "text": "",
                "feedback": "",
                "enhanced_version": "",
                "issues": [],
                "error": "❌ Invalid file (empty or >200MB)",
            }

        # Normalize doc_type
        doc_type = (doc_type or "").strip().lower()
        if doc_type == "resume":
            doc_type = "cv"

        # Detect file type and ensure supported
        file_type = self._detect_file_type(file_bytes, filename)
        logging.info(f"Detected file type: {file_type} for {filename}")

        if not self._is_supported(file_type, doc_type):
            return {
                "text": "",
                "feedback": "",
                "enhanced_version": "",
                "issues": [],
                "error": f"❌ Unsupported {file_type} format for {doc_type} analysis",
            }

        # ---- Extract text (prefer the shared helper where applicable) ----
        if file_type == "pdf":
            text = parse_uploaded_file(file_bytes, "application/pdf")
        elif file_type in {"jpg", "png"}:
            text = parse_uploaded_file(file_bytes, f"image/{file_type}")
        elif file_type == "docx":
            # native DOCX extraction
            import io
            from docx import Document
            text = "\n".join(
                Document(io.BytesIO(file_bytes)).paragraphs[i].text
                for i in range(len(Document(io.BytesIO(file_bytes)).paragraphs))
            )
        elif file_type == "txt":
            text = file_bytes.decode("utf-8", errors="ignore")
        else:
            return {
                "text": "",
                "feedback": "",
                "enhanced_version": "",
                "issues": [],
                "error": f"❌ Unsupported file type: {file_type}",
            }

        if not text or not text.strip():
            return {
                "text": "",
                "feedback": "",
                "enhanced_version": "",
                "issues": [],
                "error": "❌ Extracted text is empty",
            }

        # ---- Generate LLM analysis (purpose-aware) ----
        analysis = self._generate_analysis(text, doc_type, purpose=purpose, extra_context=extra_context)

        return {
            "text": analysis.get("text", text),
            "feedback": analysis.get("feedback", ""),
            # keep enhanced_version available for any doc_type (front-end decides showing)
            "enhanced_version": analysis.get("enhanced_version", ""),
            "issues": analysis.get("issues", []),
            "error": analysis.get("error", ""),
        }

    except Exception as e:
        logging.exception(f"Document analysis failed for {filename}: {str(e)}")
        return {
            "text": "",
            "feedback": "",
            "enhanced_version": "",
            "issues": [],
            "error": f"❌ Analysis failed: {str(e)}",
        }

    def _generate_analysis(self, text: str, doc_type: str, purpose: Optional[str] = None, extra_context: Optional[str] = None) -> Dict:
    """
    Purpose-aware reviewer. Returns JSON:
      {
        "feedback": str,
        "enhanced_version": str,
        "issues": [{"excerpt": str, "issue": str, "suggested_fix": str}, ...]
      }
    """
    try:
        # conservative input length to avoid model overrun
        # give SOP/CV more room; others slightly less
        max_chars = 4000 if doc_type in {"sop", "cv", "motivation letter"} else 2000
        trimmed = (text[:max_chars].rstrip() + "...") if len(text) > max_chars else text

        # Normalize controls
        purpose_norm = (purpose or "").strip()
        extra = (extra_context or "").strip()

        # Build a crisp system message
        # The model must return strict JSON with fields we expect.
        system_msg = (
            "You are a precise, purpose-aware document reviewer. "
            "Always return STRICT JSON with keys: "
            "feedback (string), enhanced_version (string), issues (array of objects with keys excerpt, issue, suggested_fix). "
            "Do NOT include markdown, backticks, or commentary outside JSON. "
            "Your job: "
            "1) Give concise, actionable feedback tailored to the stated purpose. "
            "2) Produce a professionally rewritten version (keep content truthful; do not fabricate credentials). "
            "3) Identify specific problems as an array of issues, each with a short excerpt (from the user's text), the issue, and a suggested fix."
        )

        # Purpose instructions
        # We keep them compact and targeted.
        purpose_hint = ""
        if purpose_norm:
            if purpose_norm.lower().startswith("masters"):
                purpose_hint = (
                    "Target: Master's admission. Emphasize clarity of goals, relevant coursework/projects, "
                    "impact, and fit with target program; avoid generic statements."
                )
            elif purpose_norm.lower().startswith("phd"):
                purpose_hint = (
                    "Target: PhD admission. Prioritize research readiness, problem statements, methodology, "
                    "publications, alignment with advisor/lab, and long-term research vision."
                )
            elif purpose_norm.lower().startswith("job"):
                purpose_hint = (
                    "Target: Job application. Emphasize quantified achievements, skills relevant to the role, "
                    "ATS-friendly phrasing, and clear impact aligned to the job description."
                )
            elif purpose_norm.lower().startswith("email"):
                purpose_hint = (
                    "Target: Email to professor. Keep it brief and respectful; show genuine alignment with their research, "
                    "mention 1–2 relevant works, and ask a specific question or next step."
                )
            else:
                purpose_hint = f"Target: {purpose_norm}. Keep tone and structure appropriate to this purpose."

        # Extra context (e.g., JD bullets, program link themes)
        extra_hint = f"Additional context (may guide tailoring): {extra}" if extra else "No extra context provided."

        user_payload = {
            "document_type": doc_type,
            "purpose": purpose_norm or "",
            "extra_context": extra or "",
            "text": trimmed,
            "instructions": [
                purpose_hint,
                "Make feedback concise and actionable (3–8 bullet points).",
                "Enhanced version must read naturally and reflect the user's content; do not invent facts.",
                "Issues array: 3–10 items with short excerpt, issue, and suggested_fix.",
            ],
        }

        completion = self.client.chat.completions.create(
            model="gpt-4-turbo",
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": system_msg},
                {"role": "user", "content": json.dumps(user_payload, ensure_ascii=False)}
            ],
            temperature=0.3,
            max_tokens=1400,
        )

        raw = completion.choices[0].message.content or "{}"
        data = {}
        try:
            data = json.loads(raw)
        except Exception:
            data = {
                "feedback": "Could not parse structured output.",
                "enhanced_version": "",
                "issues": []
            }

        self.metrics["gpt_calls"] += 1

        # Normalize issues to a list of dicts
        issues = data.get("issues", [])
        if not isinstance(issues, list):
            issues = []
        normalized_issues = []
        for it in issues:
            normalized_issues.append({
                "excerpt": (it or {}).get("excerpt", "")[:400],
                "issue": (it or {}).get("issue", "")[:400],
                "suggested_fix": (it or {}).get("suggested_fix", "")[:400],
            })

        return {
            "text": trimmed,
            "feedback": data.get("feedback", ""),
            "enhanced_version": data.get("enhanced_version", ""),
            "issues": normalized_issues,
            "error": "",
        }

    except Exception as e:
        logging.exception("Error in _generate_analysis")
        return {
            "text": text,
            "feedback": "Analysis failed.",
            "enhanced_version": "",
            "issues": [],
            "error": str(e),
        }


    # ---------- Chat responses with live scholarships (left intact) ----------
    def generate_response(self, prompt: str, context: List[Dict] = None) -> str:
        cache_key = hashlib.md5(prompt.encode()).hexdigest()
        if cache_key in self.cache:
            self.metrics["cache_hits"] += 1
            return self.cache[cache_key]

        start_time = datetime.now()

        try:
            scholarships = self._fetch_scholarships()
            scholarship_text = "\n".join([f"- {item['title']} ({item['link']})" for item in scholarships])
            enhanced_prompt = f"""
You are a scholarship assistant. Here is a list of current scholarships fetched live:

{scholarship_text}

Now, answer this user query using the above info where possible:
{prompt}
"""
            response = self.client.chat.completions.create(
                model="gpt-4-turbo",
                messages=context or [{"role": "user", "content": enhanced_prompt}],
                temperature=0.7,
                max_tokens=1500,
            )
            response_time = (datetime.now() - start_time).total_seconds()
            self.metrics["response_times"].append(response_time)
            self.metrics["gpt_calls"] += 1

            result = response.choices[0].message.content
            self.cache.set(cache_key, result, expire=24 * 3600)
            return result
        except Exception as e:
            logging.error(f"Error generating response: {str(e)}")
            return "Sorry, I encountered an error processing your request. Please try again."

    # ---------- Scholarships / Universities (unchanged behavior) ----------
    def find_universities(self) -> Dict:
        try:
            local_results = self._query_local_db()
            if local_results:
                return {"source": "local_db", "data": local_results}
            api_results = self._query_university_api()
            if api_results:
                return {"source": "api", "data": api_results}
            gpt_results = self._generate_gpt_recommendations()
            return {"source": "gpt", "data": gpt_results}
        except Exception as e:
            logging.error(f"University search failed: {str(e)}")
            return {"source": "error", "data": "Unable to fetch university information at this time"}

    def _query_local_db(self) -> Optional[List]:
        return [
            uni
            for uni in self.services["university_db"]
            if uni["country"] in self.user["academic"]["target_countries"]
            and (self.user["academic"]["field"] or "").lower() in uni.get("programs", [])
        ]

    def find_scholarships(self, query: str = None) -> List[Dict]:
        all_scholarships = self._fetch_scholarships()
        if query:
            query_embed = self.embedding_model.encode(query)
            scholarships_with_scores = []
            for scholarship in all_scholarships:
                title_embed = self.embedding_model.encode(scholarship["title"])
                score = cosine_similarity([query_embed], [title_embed])[0][0]
                scholarships_with_scores.append((scholarship, score))
            return [x[0] for x in sorted(scholarships_with_scores, key=lambda x: x[1], reverse=True)[:5]]
        return all_scholarships[:10]

    def suggest_next_steps(self) -> List[str]:
        context = f"""
User Profile:
- Degree: {self.user['academic']['degree']}
- Field: {self.user['academic']['field']}
- Target Countries: {self.user['academic']['target_countries']}
Documents Uploaded: {len(self.user['documents'].keys())}
"""
        return self.generate_response("Suggest 3 personalized next steps for this student:\n" + context).split("\n")

    def validate_input(self, text: str) -> bool:
        blacklist = ["ignore previous", "###", "system prompt", "as an AI", "your instructions"]
        return not any(phrase in text.lower() for phrase in blacklist)

    def get_metrics(self) -> Dict:
        try:
            avg_time = 0.0
            if self.metrics["response_times"]:
                avg_time = float(np.mean(self.metrics["response_times"]))
            return {
                "gpt_calls": self.metrics["gpt_calls"],
                "cache_hit_rate": f"{(self.metrics['cache_hits']/max(1, self.metrics['gpt_calls']))*100:.2f}%",
                "avg_response_time": f"{avg_time:.2f}s",
                "total_queries": self.user["engagement"]["query_count"],
            }
        except Exception as e:
            logging.error(f"Metrics generation failed: {str(e)}")
            return {"error": "Metrics temporarily unavailable", "details": str(e)}

    def _fetch_scholarships(self) -> List[Dict]:
        all_scholarships = []
        for feed_url in self.services["scholarship_feeds"]:
            try:
                feed = feedparser.parse(feed_url)
                all_scholarships.extend(
                    [
                        {
                            "title": entry.title,
                            "link": entry.link,
                            "summary": entry.get("summary", "No details available"),
                        }
                        for entry in feed.entries[:5]
                    ]
                )
            except Exception as e:
                logging.error(f"Failed to parse {feed_url}: {str(e)}")
        return all_scholarships

    # ----- Stubs left intentionally as in your original code -----
    def _query_university_api(self) -> List[Dict]:
        pass

    def _generate_gpt_recommendations(self) -> List[Dict]:
        pass

